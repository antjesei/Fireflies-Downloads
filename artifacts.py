"""Rendert Transcript/Summary als MD und DOCX. Kein Fireflies-Branding."""
from __future__ import annotations

import re
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from fireflies_api import Summary, Transcript
from postprocess import ProcessedBlock, format_timestamp


# ---------- Mini-Markdown-Renderer für DOCX ----------

# Bullet-Marker am Zeilenanfang (mit optionaler Einrückung):
# -, *, •, →, →
_BULLET_RE = re.compile(r"^(\s*)([-*•→▪▶►]|→|•)\s+(.*)$")
_BOLD_ITALIC_RE = re.compile(r"\*\*(.+?)\*\*|__(.+?)__|\*(.+?)\*|_(.+?)_")


def _add_inline_runs(paragraph, text: str) -> None:
    """Fügt Text mit **fett** / *kursiv* / __fett__ / _kursiv_ als Runs in den Paragraph ein."""
    pos = 0
    for m in _BOLD_ITALIC_RE.finditer(text):
        # Plain-Text vor dem Match
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        bold_g = m.group(1) or m.group(2)
        italic_g = m.group(3) or m.group(4)
        run = paragraph.add_run(bold_g if bold_g else italic_g)
        if bold_g:
            run.bold = True
        else:
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_markdown_block(doc: Document, body: str, base_heading_level: int = 2) -> None:
    """Rendert einen Markdown-artigen Body in den Doc.

    Unterstützt:
    - `#`, `##`, `###` Headings (verschoben relativ zu base_heading_level)
    - Bullet-Listen mit `-`, `*`, `•`, `→`
    - Inline `**fett**`, `*kursiv*`
    - Leerzeilen als Absatztrenner
    """
    for raw in body.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue

        # Headings
        m_h = re.match(r"^(#{1,6})\s+(.*)$", line.lstrip())
        if m_h:
            depth = len(m_h.group(1))
            text = m_h.group(2).strip()
            # **fett** in Heading-Text bereinigen (Heading ist ohnehin fett)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"__(.+?)__", r"\1", text)
            level = min(base_heading_level + depth - 1, 6)
            doc.add_heading(text, level=level)
            continue

        # Bullet-Listen
        m_b = _BULLET_RE.match(line)
        if m_b:
            indent_spaces = len(m_b.group(1))
            content = m_b.group(3).strip()
            # Doppelte Marker am Inhaltsanfang entfernen (z.B. "- → text" -> "text")
            content = re.sub(r"^([-*•→▪▶►]\s+)+", "", content)
            # python-docx kennt "List Bullet", "List Bullet 2", "List Bullet 3"
            level = min(indent_spaces // 2 + 1, 3)
            style = "List Bullet" if level == 1 else f"List Bullet {level}"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, content)
            continue

        # Normaler Absatz
        p = doc.add_paragraph()
        _add_inline_runs(p, line)


def _add_page_numbers(doc: Document) -> None:
    """Fügt eine zentrierte Seitenzahl im Footer aller Sections ein.

    Format: 'Seite X von Y' — verwendet Word-Felder PAGE und NUMPAGES.
    """
    for section in doc.sections:
        footer = section.footer
        # Erster (und meist einziger) Paragraph im Footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = 1  # CENTER
        # Vorhandenen Text leeren
        for run in list(p.runs):
            run.text = ""

        # "Seite "
        run_prefix = p.add_run("Seite ")
        run_prefix.font.size = Pt(9)
        run_prefix.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # PAGE-Feld
        _add_field(p, "PAGE")

        # " von "
        run_mid = p.add_run(" von ")
        run_mid.font.size = Pt(9)
        run_mid.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # NUMPAGES-Feld
        _add_field(p, "NUMPAGES")


def _add_field(paragraph, instr: str) -> None:
    """Hilfsfunktion: fügt ein Word-Feld (z.B. PAGE, NUMPAGES) in einen Paragraph ein."""
    run = paragraph.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)


def _set_calibri(doc: Document) -> None:
    """Setzt Calibri als Standardschrift für das gesamte Dokument."""
    styles = doc.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(11)
    # Auch für asiatische / komplexe Schriften setzen
    rPr = styles.element.get_or_add_rPr()
    for tag in (qn("w:rFonts"),):
        el = rPr.find(tag)
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:ascii"), "Calibri")
        el.set(qn("w:hAnsi"), "Calibri")
        el.set(qn("w:cs"), "Calibri")


def _meta_lines(t: Transcript) -> list[str]:
    date_str = t.date.strftime("%Y-%m-%d %H:%M")
    dur_min = int(round(t.duration))
    participants = ", ".join(t.participants) if t.participants else "-"
    return [
        f"Titel: {t.title}",
        f"Datum: {date_str}",
        f"Dauer: {dur_min} Min.",
        f"Teilnehmer: {participants}",
    ]


# ---------- Transcript Markdown ----------

def render_transcript_md(t: Transcript, blocks: list[ProcessedBlock], path: str) -> None:
    lines: list[str] = [f"# {t.title}", ""]
    for meta in _meta_lines(t):
        lines.append(meta)
    lines.append("")
    lines.append("---")
    lines.append("")

    for b in blocks:
        ts = format_timestamp(b.start_time)
        if b.speaker_name:
            lines.append(f"[{ts}] **{b.speaker_name}**: {b.text}")
        else:
            lines.append(f"[{ts}] {b.text}")
        lines.append("")  # Leerzeile = Absatz-Trenner in MD

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines).rstrip() + "\n")


# ---------- Transcript DOCX ----------

def render_transcript_docx(t: Transcript, blocks: list[ProcessedBlock], path: str) -> None:
    doc = Document()
    _set_calibri(doc)
    _add_page_numbers(doc)

    # Titel
    h = doc.add_heading(t.title, level=1)
    # Metadaten
    for meta in _meta_lines(t):
        p = doc.add_paragraph(meta)
        for run in p.runs:
            run.font.size = Pt(10)
    doc.add_paragraph()  # Spacer

    # Sätze
    for b in blocks:
        ts = format_timestamp(b.start_time)
        p = doc.add_paragraph()
        # Timestamp — grau, klein
        ts_run = p.add_run(f"[{ts}] ")
        ts_run.font.size = Pt(9)
        ts_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        if b.speaker_name:
            sp_run = p.add_run(f"{b.speaker_name}: ")
            sp_run.bold = True

        p.add_run(b.text)

    doc.save(path)


# ---------- Summary rendering ----------

def _summary_sections(s: Summary) -> list[tuple[str, str]]:
    """Gibt (Heading, Body)-Tupel zurück, nur für Sections mit Inhalt.

    Reihenfolge: Erst die ausführlichen Inhalte (nach UI-Klick "Refine Summary"),
    danach die Standard-Sektionen.
    """
    sections: list[tuple[str, str]] = []

    # Ausführliche Notizen (Refine Summary / Get more details)
    if s.notes and s.notes.strip():
        sections.append(("Ausführliche Notizen", s.notes.strip()))

    # Standard-Sections
    if s.overview and s.overview.strip():
        sections.append(("Overview", s.overview.strip()))
    if s.short_summary and s.short_summary.strip():
        sections.append(("Kurzfassung", s.short_summary.strip()))
    if s.short_overview and s.short_overview.strip() and s.short_overview.strip() != (s.overview or "").strip():
        sections.append(("Kurzüberblick", s.short_overview.strip()))
    if s.gist and s.gist.strip():
        sections.append(("Gist", s.gist.strip()))
    if s.action_items and s.action_items.strip():
        sections.append(("Action Items", s.action_items.strip()))
    if s.bullet_gist and s.bullet_gist.strip():
        sections.append(("Kernpunkte", s.bullet_gist.strip()))
    if s.shorthand_bullet and s.shorthand_bullet.strip():
        sections.append(("Shorthand", s.shorthand_bullet.strip()))
    if s.topics_discussed:
        sections.append(("Themen", "\n".join(f"- {x}" for x in s.topics_discussed)))
    if s.keywords:
        sections.append(("Keywords", ", ".join(s.keywords)))
    return sections


def render_summary_md(t: Transcript, path: str) -> None:
    lines: list[str] = [f"# {t.title} — Summary", ""]
    for meta in _meta_lines(t):
        lines.append(meta)
    lines.append("")
    lines.append("---")
    lines.append("")
    for heading, body in _summary_sections(t.summary):
        lines.append(f"## {heading}")
        lines.append("")
        lines.append(body)
        lines.append("")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines).rstrip() + "\n")


def render_summary_docx(t: Transcript, path: str) -> None:
    doc = Document()
    _set_calibri(doc)
    _add_page_numbers(doc)
    doc.add_heading(f"{t.title} — Summary", level=1)
    for meta in _meta_lines(t):
        p = doc.add_paragraph(meta)
        for run in p.runs:
            run.font.size = Pt(10)
    doc.add_paragraph()
    for heading, body in _summary_sections(t.summary):
        doc.add_heading(heading, level=2)
        # Body als Mini-Markdown rendern (Headings, Bullets, **fett**, *kursiv*)
        _add_markdown_block(doc, body, base_heading_level=3)
    doc.save(path)
